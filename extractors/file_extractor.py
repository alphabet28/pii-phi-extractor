"""
Stage 1: File Extraction
Extracts text and metadata from multiple file formats.
Returns normalized extraction results for each document.
"""

import os
import re
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict, field
from email import policy
from email.parser import BytesParser
from bs4 import BeautifulSoup

# Third-party imports
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import extract_msg
except ImportError:
    extract_msg = None


logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Normalized extraction result for all file formats."""
    filename: str
    filepath: str
    filetype: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_status: str = "success"  # "success" | "error" | "empty" | "needs_ocr"
    error: Optional[str] = None
    attachment_extractions: List[Dict] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return asdict(self)


class FileExtractor:
    """
    Orchestrates text extraction from various file formats.
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.txt', '.md', '.docx', '.xlsx', '.pdf', '.eml', '.msg',
        '.doc', '.xls'  # older formats (basic support via error handling)
    }

    # Attachment extraction settings
    MAX_ATTACHMENT_DEPTH = 3
    ATTACHMENT_SIZE_LIMIT = 50 * 1024 * 1024  # 50 MB

    def __init__(self, log_level: str = "INFO"):
        """Initialize the file extractor."""
        logging.basicConfig(level=log_level)
        self.logger = logger

    def extract(self, filepath: str) -> ExtractionResult:
        """
        Main entry point: extract text from a file.
        Routes to appropriate handler based on file extension.
        
        Args:
            filepath: Absolute or relative path to the file
            
        Returns:
            ExtractionResult with normalized output
        """
        filepath = str(filepath)
        path_obj = Path(filepath)

        # Validate file exists
        if not os.path.exists(filepath):
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype=path_obj.suffix.lower(),
                text="",
                extraction_status="error",
                error=f"File not found: {filepath}"
            )

        # Get file extension
        ext = path_obj.suffix.lower()

        # Dispatch to appropriate handler
        try:
            if ext in {'.txt', '.md'}:
                return self._extract_text(filepath)
            elif ext == '.docx':
                return self._extract_docx(filepath)
            elif ext == '.xlsx':
                return self._extract_xlsx(filepath)
            elif ext == '.pdf':
                return self._extract_pdf(filepath)
            elif ext == '.eml':
                return self._extract_eml(filepath)
            elif ext == '.msg':
                return self._extract_msg(filepath)
            elif ext in {'.doc', '.xls'}:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype=ext,
                    text="",
                    extraction_status="error",
                    error=f"Unsupported format (legacy Office): {ext}. Please convert to .docx/.xlsx"
                )
            else:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype=ext,
                    text="",
                    extraction_status="error",
                    error=f"Unsupported file type: {ext}"
                )
        except Exception as e:
            self.logger.error(f"Error extracting {filepath}: {e}")
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype=ext,
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _extract_text(self, filepath: str) -> ExtractionResult:
        """Extract from plain text files (.txt, .md)."""
        path_obj = Path(filepath)
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read().strip()

            if not text:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype=path_obj.suffix.lower(),
                    text="",
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype=path_obj.suffix.lower(),
                text=text,
                extraction_status="success"
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype=path_obj.suffix.lower(),
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _extract_docx(self, filepath: str) -> ExtractionResult:
        """Extract from Word documents (.docx)."""
        path_obj = Path(filepath)

        if DocxDocument is None:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.docx',
                text="",
                extraction_status="error",
                error="python-docx not installed"
            )

        try:
            doc = DocxDocument(filepath)
            texts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)

            full_text = "\n".join(texts).strip()

            if not full_text:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.docx',
                    text="",
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.docx',
                text=full_text,
                metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
                extraction_status="success"
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.docx',
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _extract_xlsx(self, filepath: str) -> ExtractionResult:
        """Extract from Excel spreadsheets (.xlsx)."""
        path_obj = Path(filepath)

        if load_workbook is None:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.xlsx',
                text="",
                extraction_status="error",
                error="openpyxl not installed"
            )

        try:
            wb = load_workbook(filepath, data_only=True)
            texts = []
            sheet_names = []

            for sheet in wb.sheetnames:
                sheet_names.append(sheet)
                ws = wb[sheet]
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            texts.append(str(cell.value))

            full_text = "\n".join(texts).strip()

            if not full_text:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.xlsx',
                    text="",
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.xlsx',
                text=full_text,
                metadata={"sheets": sheet_names, "sheet_count": len(sheet_names)},
                extraction_status="success"
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.xlsx',
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _extract_pdf(self, filepath: str) -> ExtractionResult:
        """
        Extract from PDF files.
        Detects scanned PDFs (empty text) and flags for OCR.
        """
        path_obj = Path(filepath)

        if pdfplumber is None:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.pdf',
                text="",
                extraction_status="error",
                error="pdfplumber not installed"
            )

        try:
            texts = []
            scanned_pages = 0

            with pdfplumber.open(filepath) as pdf:
                page_count = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    text = text.strip()

                    if not text:
                        scanned_pages += 1
                    else:
                        texts.append(text)

                    # Also extract tables from this page
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                for cell in row:
                                    if cell:
                                        texts.append(str(cell))

            full_text = "\n".join(texts).strip()

            # All pages are scanned (no extractable text)
            if scanned_pages == page_count:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.pdf',
                    text="",
                    metadata={"page_count": page_count, "scanned_pages": scanned_pages},
                    extraction_status="needs_ocr",
                    error="PDF is scanned (image-based); requires OCR processing"
                )

            # Some pages are scanned but we got some text
            if scanned_pages > 0 and full_text:
                self.logger.warning(
                    f"{filepath}: {scanned_pages}/{page_count} pages are scanned; "
                    f"extracted {len(full_text)} chars from remaining pages"
                )

            if not full_text:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.pdf',
                    text="",
                    metadata={"page_count": page_count, "scanned_pages": scanned_pages},
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.pdf',
                text=full_text,
                metadata={"page_count": page_count, "scanned_pages": scanned_pages},
                extraction_status="success"
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.pdf',
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _extract_eml(self, filepath: str) -> ExtractionResult:
        """
        Extract from email files (.eml).
        Parses headers, body, and attachments recursively.
        """
        path_obj = Path(filepath)

        try:
            with open(filepath, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)

            parts = []
            attachment_extractions = []

            # Extract headers
            headers = {
                'From': msg.get('From', ''),
                'To': msg.get('To', ''),
                'Cc': msg.get('Cc', ''),
                'Bcc': msg.get('Bcc', ''),
                'Subject': msg.get('Subject', ''),
                'Date': msg.get('Date', '')
            }

            for key, value in headers.items():
                if value:
                    parts.append(f"{key}: {value}")

            # Walk MIME tree
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = part.get("Content-Disposition", "")

                    # Extract text/plain body
                    if content_type == "text/plain" and "attachment" not in content_disposition:
                        payload = part.get_payload(decode=True)
                        if payload:
                            text = payload.decode('utf-8', errors='replace')
                            if text.strip():
                                parts.append(text)

                    # Extract text/html body
                    elif content_type == "text/html" and "attachment" not in content_disposition:
                        payload = part.get_payload(decode=True)
                        if payload:
                            html_text = payload.decode('utf-8', errors='replace')
                            # Strip HTML tags
                            soup = BeautifulSoup(html_text, 'html.parser')
                            text = soup.get_text()
                            if text.strip():
                                parts.append(text)

                    # Handle attachments
                    elif "attachment" in content_disposition:
                        filename = part.get_filename()
                        if filename:
                            attachment_extractions.extend(
                                self._process_email_attachment(part, filepath, depth=0)
                            )
            else:
                # Non-multipart message
                payload = msg.get_payload(decode=True)
                if payload:
                    text = payload.decode('utf-8', errors='replace')
                    if text.strip():
                        parts.append(text)

            full_text = "\n".join(parts).strip()

            if not full_text and not attachment_extractions:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.eml',
                    text="",
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.eml',
                text=full_text,
                metadata={"headers": headers},
                extraction_status="success",
                attachment_extractions=attachment_extractions
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.eml',
                text="",
                extraction_status="error",
                error=str(e)
            )

    def _process_email_attachment(
        self, part: Any, parent_filepath: str, depth: int
    ) -> List[Dict[str, Any]]:
        """
        Recursively process email attachments.
        Extracts supported file types and re-enters the pipeline.
        """
        results = []

        if depth >= self.MAX_ATTACHMENT_DEPTH:
            self.logger.warning(f"Max attachment depth reached for {parent_filepath}")
            return results

        filename = part.get_filename()
        if not filename:
            return results

        # Check file extension
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return results

        # Check size
        try:
            payload = part.get_payload(decode=True)
            if payload and len(payload) > self.ATTACHMENT_SIZE_LIMIT:
                self.logger.warning(
                    f"Attachment {filename} in {parent_filepath} exceeds size limit"
                )
                return results
        except Exception as e:
            self.logger.error(f"Error processing attachment {filename}: {e}")
            return results

        # Write to temp file and extract
        try:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(payload)
                tmp_path = tmp.name

            # Recursively extract
            attachment_result = self.extract(tmp_path)
            
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except:
                pass

            # Add attachment metadata
            attachment_dict = attachment_result.to_dict()
            attachment_dict['original_filename'] = filename
            attachment_dict['parent_filepath'] = parent_filepath
            results.append(attachment_dict)

        except Exception as e:
            self.logger.error(f"Error extracting attachment {filename}: {e}")

        return results

    def _extract_msg(self, filepath: str) -> ExtractionResult:
        """
        Extract from Outlook MSG files.
        Falls back to error if extract_msg is not available.
        """
        path_obj = Path(filepath)

        if extract_msg is None:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.msg',
                text="",
                extraction_status="error",
                error="extract-msg not installed. Install with: pip install extract-msg"
            )

        try:
            msg = extract_msg.Message(filepath)
            parts = []

            # Extract headers
            if msg.sender:
                parts.append(f"From: {msg.sender}")
            if msg.to:
                parts.append(f"To: {msg.to}")
            if msg.cc:
                parts.append(f"Cc: {msg.cc}")
            if msg.subject:
                parts.append(f"Subject: {msg.subject}")
            if msg.date:
                parts.append(f"Date: {msg.date}")

            # Extract body
            if msg.body:
                parts.append(msg.body)

            full_text = "\n".join(parts).strip()

            if not full_text:
                return ExtractionResult(
                    filename=path_obj.name,
                    filepath=filepath,
                    filetype='.msg',
                    text="",
                    extraction_status="empty"
                )

            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.msg',
                text=full_text,
                metadata={
                    "sender": msg.sender,
                    "to": msg.to,
                    "subject": msg.subject
                },
                extraction_status="success"
            )
        except Exception as e:
            return ExtractionResult(
                filename=path_obj.name,
                filepath=filepath,
                filetype='.msg',
                text="",
                extraction_status="error",
                error=str(e)
            )
